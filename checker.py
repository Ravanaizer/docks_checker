import re
from dataclasses import dataclass
from enum import Enum
from typing import List

from docx import Document
from docx.oxml.ns import qn


class Severity(Enum):
    CRITICAL = "CRITICAL"
    WARNING = "WARNING"
    INFO = "INFO"


@dataclass
class ValidationError:
    rule: str
    message: str
    severity: Severity
    location: str = ""


class DocumentArchitectureValidator:
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.errors: List[ValidationError] = []

        # Clean empty paragraphs
        self._clean_empty_paragraphs()

        self.paragraphs = self.doc.paragraphs
        self.tables = self.doc.tables

        # Find boundary between main text and appendices
        self.main_doc_end_idx = self._find_main_document_boundary()

        # Split parts
        self.main_paragraphs = self.paragraphs[: self.main_doc_end_idx]
        self.appendix_paragraphs = self.paragraphs[self.main_doc_end_idx :]

        # Normalize text
        self.main_text = self._normalize_text(
            "\n".join([p.text.strip() for p in self.main_paragraphs if p.text.strip()])
        )
        self.full_text = self._normalize_text(
            "\n".join([p.text.strip() for p in self.paragraphs if p.text.strip()])
        )

    def validate(self) -> List[ValidationError]:
        self._check_heading()
        self._check_body_font_size()
        self._check_preamble_structure()
        self._check_command_word()
        self._check_control_clause_position()
        self._check_signature_block()
        self._check_appendix_format()
        self._check_table_fonts()
        self._check_reduction_position()
        return self.errors

    def _clean_empty_paragraphs(self):
        """Remove paragraphs that contain only whitespace."""
        for p in reversed(self.doc.paragraphs):
            if not p.text.strip():
                p._element.getparent().remove(p._element)

    def _normalize_text(self, text: str) -> str:
        """Replace multiple whitespaces/tabs with a single space."""
        return re.sub(r"\s+", " ", text).strip()

    def _get_effective_font_size(self, run) -> float:
        """Get actual font size considering style inheritance."""
        if run.font.size is not None:
            return run.font.size.pt

        # Check paragraph style
        paragraph = run._parent
        if paragraph and paragraph.style:
            try:
                style_font = paragraph.style.font
                if style_font and style_font.size:
                    return style_font.size.pt
            except Exception:
                pass

        # Check document default style
        try:
            doc_style = self.doc.style
            if doc_style and doc_style.font and doc_style.font.size:
                return doc_style.font.size.pt
        except Exception:
            pass

        return 12.0

    def _find_main_document_boundary(self) -> int:
        """Find where main text ends and appendices begin."""
        appendix_markers = [r"^приложение\s*[№№]?\s*[а-яё1-9]", r"^прил\.\s*"]
        para_idx = 0

        for child in self.doc.element.body:
            if child.tag == qn("w:p"):
                if para_idx < len(self.paragraphs):
                    para = self.paragraphs[para_idx]
                    text = self._normalize_text(para.text.strip().lower())
                    if any(re.search(pattern, text) for pattern in appendix_markers):
                        return para_idx
                    para_idx += 1

            elif child.tag == qn("w:tbl"):
                for row in child.iter(qn("w:tr")):
                    for cell in row.iter(qn("w:tc")):
                        for p in cell.iter(qn("w:p")):
                            cell_text = " ".join(p.itertext()).strip().lower()
                            cell_text = self._normalize_text(cell_text)
                            if any(
                                re.search(pattern, cell_text)
                                for pattern in appendix_markers
                            ):
                                return para_idx
        return len(self.paragraphs)

    def _check_heading(self):
        patterns = [
            r"об\s+утверждении",
            r"о\s+подготовке\s+номенклатуры\s+дел",
            r"о\s+внесении\s+изменений\s+в\s+приказ",
        ]
        found = False

        # Check tables at start
        for table in self.tables[:3]:
            text = "\n".join(self._get_table_text(table)).lower()
            if any(re.search(p, text) for p in patterns):
                found = True
                if not self._check_table_font_size(table, 10):
                    self.errors.append(
                        ValidationError(
                            "HEADING",
                            "Font size in heading table is incorrect (expected 10pt)",
                            Severity.CRITICAL,
                            "Table at start",
                        )
                    )
                break

        # Check paragraphs at start
        if not found:
            for para in self.paragraphs[:10]:
                text = self._normalize_text(para.text).lower()
                if any(re.search(p, text) for p in patterns):
                    found = True
                    for run in para.runs:
                        if run.text.strip():
                            size = self._get_effective_font_size(run)
                            if not (9.5 <= size <= 10.5):
                                self.errors.append(
                                    ValidationError(
                                        "HEADING",
                                        f"Font size is incorrect: {size}pt (expected 10pt)",
                                        Severity.CRITICAL,
                                        "Start of document",
                                    )
                                )
                    break

        if not found:
            self.errors.append(
                ValidationError(
                    "HEADING",
                    "Heading is missing.",
                    Severity.CRITICAL,
                    "Start of document",
                )
            )

    def _check_body_font_size(self):
        sample = self.main_paragraphs[:-2]
        for para in sample[::5]:
            for run in para.runs:
                if run.text.strip():
                    size = self._get_effective_font_size(run)
                    if not (11.5 <= size <= 12.5):
                        self.errors.append(
                            ValidationError(
                                "BODY_FONT",
                                f"Body font size is incorrect: {size}pt (expected 12pt)",
                                Severity.WARNING,
                                "Document body",
                            )
                        )
                        return

    def _check_preamble_structure(self):
        patterns = [
            r"в\s+соответствии",
            r"во\s+исполнение",
            r"согласно",
            r"в\s+целях",
            r"в\s+связи\s+с",
        ]
        text = self._normalize_text(
            "\n".join([p.text for p in self.paragraphs[:10]])
        ).lower()

        if not any(re.search(p, text) for p in patterns):
            self.errors.append(
                ValidationError(
                    "PREAMBLE",
                    "Justification (preamble) is missing.",
                    Severity.CRITICAL,
                    "Start of document",
                )
            )

    def _check_command_word(self):
        patterns = [r"п\s*р\s*и\s*к\s*а\s*з\s*ы\s*в\s*а\s*ю", r"приказываю"]
        if not any(re.search(p, self.full_text.lower()) for p in patterns):
            self.errors.append(
                ValidationError(
                    "COMMAND",
                    "Word 'Приказываю' contains extra spaces or is missing.",
                    Severity.CRITICAL,
                    "Dispositive section",
                )
            )

    def _check_control_clause_position(self):
        patterns = [
            r"контроль\s+(за\s+исполнением|оставляю|возложить)",
            r"контроль\s+оставляю\s+за\s+собой",
            r"возложить\s+контроль",
        ]
        text = self._normalize_text(
            "\n".join([p.text for p in self.main_paragraphs[-10:]])
        ).lower()

        if not any(re.search(p, text) for p in patterns):
            self.errors.append(
                ValidationError(
                    "CONTROL",
                    "Clause about control of order execution is missing.",
                    Severity.CRITICAL,
                    "End of dispositive section",
                )
            )

    def _check_reduction_position(self):
        patterns = [
            r"довести\s+настоящий\s+приказ\s+до\s+сведения",
            r"довести\s+до\s+сведения",
        ]
        text = self._normalize_text(
            "\n".join([p.text for p in self.main_paragraphs[-10:]])
        ).lower()

        if not any(re.search(p, text) for p in patterns):
            self.errors.append(
                ValidationError(
                    "REDUCTION",
                    "No section requiring information distribution.",
                    Severity.CRITICAL,
                    "End of dispositive section",
                )
            )

    def _check_signature_block(self):
        patterns = [
            r"[а-яё]+\s+[а-яё]\.[а-яё]\.",
            r"[а-яё]\.[а-яё]\.\s+[а-яё]+",
            r"директор.*[а-яё]+\s+[а-яё]\.",
            r"руководитель.*[а-яё]+\s+[а-яё]\.",
        ]
        found = False

        text = self._normalize_text(
            "\n".join([p.text for p in self.main_paragraphs[-5:]])
        ).lower()
        if any(re.search(p, text) for p in patterns):
            found = True

        if not found and self.tables:
            for table in self.tables[-3:]:
                table_text = "\n".join(self._get_table_text(table)).lower()
                if any(re.search(p, table_text) for p in patterns):
                    found = True
                    break

        if not found:
            self.errors.append(
                ValidationError(
                    "SIGNATURE",
                    "No head signature block found.",
                    Severity.CRITICAL,
                    "End of main text",
                )
            )

    def _check_appendix_format(self):
        forbidden = "ёзйочьыъ"
        headers = []

        for para in self.appendix_paragraphs[:30]:
            text = self._normalize_text(para.text.lower())
            if "приложение" in text:
                headers.append({"text": text, "source": "paragraph"})

        for table in self.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = self._normalize_text(para.text.lower())
                        if "приложение" in text:
                            headers.append({"text": text, "source": "table"})

        found_letters = []
        found_numbers = []

        for h in headers:
            text = h["text"]
            letter_match = re.search(r"приложение\s*[№№]?\s*([а-яё])", text)
            number_match = re.search(r"приложение\s*[№№]?\s*(\d+)", text)

            if letter_match:
                letter = letter_match.group(1)
                found_letters.append(letter)
                if letter in forbidden:
                    self.errors.append(
                        ValidationError(
                            "APPENDIX",
                            f"Forbidden letter in appendix: '{letter.upper()}'",
                            Severity.WARNING,
                            "Appendix header",
                        )
                    )

            if number_match:
                found_numbers.append(int(number_match.group(1)))

        if found_letters:
            self._check_sequential_numbering(found_letters, "letter")
        if found_numbers:
            self._check_sequential_numbering(found_numbers, "number")

    def _check_sequential_numbering(self, items, type="letter"):
        if len(items) < 2:
            return

        if type == "letter":
            alphabet = "абвгдежзиклмнпрстуфхцшэюя"
            indices = [alphabet.index(i) if i in alphabet else -1 for i in items]
            for i in range(1, len(indices)):
                if (
                    indices[i] != -1
                    and indices[i - 1] != -1
                    and indices[i] - indices[i - 1] > 1
                ):
                    self.errors.append(
                        ValidationError(
                            "APPENDIX",
                            f"Appendix numbering not sequential: missing letter between '{alphabet[indices[i - 1]].upper()}' and '{alphabet[indices[i]].upper()}'",
                            Severity.WARNING,
                            "Appendix section",
                        )
                    )
        elif type == "number":
            sorted_items = sorted(set(items))
            for i in range(1, len(sorted_items)):
                if sorted_items[i] - sorted_items[i - 1] > 1:
                    self.errors.append(
                        ValidationError(
                            "APPENDIX",
                            f"Appendix numbering not sequential: missing number {sorted_items[i - 1] + 1}",
                            Severity.WARNING,
                            "Appendix section",
                        )
                    )

    def _check_table_fonts(self):
        for table in self.tables:
            if self._is_table_font_too_small(table):
                self.errors.append(
                    ValidationError(
                        "TABLES",
                        "Font size is too small in table.",
                        Severity.WARNING,
                        "Tables",
                    )
                )

    def _is_table_font_too_small(self, table) -> bool:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.size and run.font.size.pt < 8:
                            return True
        return False

    def _get_table_text(self, table) -> List[str]:
        lines = []
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        lines.append(para.text.strip())
        return lines

    def _check_table_font_size(self, table, expected_size, tolerance=0.5) -> bool:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip():
                            size = self._get_effective_font_size(run)
                            if (
                                size is not None
                                and abs(size - expected_size) > tolerance
                            ):
                                return False
        return True

    def print_report(self):
        if not self.errors:
            print("Document architecture meets requirements.")
            return

        critical = [e for e in self.errors if e.severity == Severity.CRITICAL]
        warnings = [e for e in self.errors if e.severity == Severity.WARNING]

        print("\n" + "=" * 70)
        print("DOCUMENT ARCHITECTURE VALIDATION REPORT")
        print("=" * 70)
        print(f"File: {self.filepath}")
        print(f"Total findings: {len(self.errors)}\n")

        if critical:
            print("CRITICAL ERRORS:")
            for err in critical:
                print(f"  {err.severity.value} | {err.rule}")
                print(f"    {err.message}")
                print()

        if warnings:
            print("WARNINGS:")
            for err in warnings:
                print(f"  {err.severity.value} | {err.rule}")
                print(f"    {err.message}")
                print()

        print("=" * 70)


if __name__ == "__main__":
    validator = DocumentArchitectureValidator("standart.docx")
    validator.validate()
    validator.print_report()
