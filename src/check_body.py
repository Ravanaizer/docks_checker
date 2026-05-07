import re

from docx.enum.section import WD_ORIENTATION

from config import Severity, ValidationError
from tables import _check_table_font_size, _get_table_text
from text_normalization import _normalize_text
from utils import _get_effective_font_size


def _check_heading(document):
    patterns = [
        r"\bоб\s+утверждении\b",
        r"\bо\s+подготовке\s+номенклатуры\s+дел\b",
        r"\bо\s+внесении\s+изменений\s+в\s+приказ\b",
    ]
    found = False

    for table in document.tables[:3]:
        cell_texts = [
            cell.text.strip()
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]
        table_text = " | ".join(cell_texts).lower()

        if not table_text:
            continue

        if any(re.search(p, table_text) for p in patterns):
            found = True
            if not _check_table_font_size(document, table, 10):
                document.errors.append(
                    ValidationError(
                        "HEADING",
                        "Font size in heading table is incorrect (expected 10pt)",
                        Severity.WARNING,
                        "Table at start",
                    )
                )
            break

    if not found:
        document.errors.append(
            ValidationError(
                "HEADING",
                "Heading is missing.",
                Severity.CRITICAL,
                "Start of document",
            )
        )

    if not found:
        for para in document.paragraphs[:10]:
            text = _normalize_text(para.text).lower()
            if any(re.search(p, text) for p in patterns):
                found = True
                for run in para.runs:
                    if run.text.strip():
                        size = _get_effective_font_size(document, run)
                        if not (9.5 <= size <= 10.5):
                            document.errors.append(
                                ValidationError(
                                    "HEADING",
                                    f"Font size is incorrect: {size}pt (expected 10pt)",
                                    Severity.CRITICAL,
                                    "Start of document",
                                )
                            )
                break

    if not found:
        document.errors.append(
            ValidationError(
                "HEADING",
                "Heading is missing.",
                Severity.CRITICAL,
                "Start of document",
            )
        )


def _check_preamble_structure(document):
    patterns = [
        r"в\s+соответствии",
        r"во\s+исполнение",
        r"согласно",
        r"в\s+целях",
        r"в\s+связи\s+с",
    ]
    text = _normalize_text(
        "\n".join([p.text for p in document.paragraphs[:10]])
    ).lower()
    if not any(re.search(p, text) for p in patterns):
        document.errors.append(
            ValidationError(
                "PREAMBLE",
                "Justification (preamble) is missing.",
                Severity.CRITICAL,
                "Start of document",
            )
        )


def _check_command_word(document):
    patterns = [r"п\sр\sи\sк\sа\sз\sы\sв\sа\s*ю", r"приказываю"]
    if not any(re.search(p, document.full_text.lower()) for p in patterns):
        document.errors.append(
            ValidationError(
                "COMMAND",
                "Word 'Приказываю' contains extra spaces or is missing.",
                Severity.CRITICAL,
                "Dispositive section",
            )
        )


def _check_control_clause_position(document):
    patterns = [
        r"контроль\s+(за\s+исполнением|оставляю|возложить)",
        r"контроль\s+оставляю\s+за\s+собой",
        r"возложить\s+контроль",
    ]
    text = _normalize_text(
        "\n".join([p.text for p in document.main_paragraphs[-10:]])
    ).lower()
    if not any(re.search(p, text) for p in patterns):
        document.errors.append(
            ValidationError(
                "CONTROL",
                "Clause about control of order execution is missing.",
                Severity.CRITICAL,
                "End of dispositive section",
            )
        )


def _check_reduction_position(document):
    patterns = [
        r"довести\s+настоящий\s+приказ\s+до\s+сведения",
        r"довести\s+до\s+сведения",
    ]
    text = _normalize_text(
        "\n".join([p.text for p in document.main_paragraphs[-10:]])
    ).lower()
    if not any(re.search(p, text) for p in patterns):
        document.errors.append(
            ValidationError(
                "REDUCTION",
                "No section requiring information distribution.",
                Severity.CRITICAL,
                "End of dispositive section",
            )
        )


def _check_signature_block(document):

    patterns = [
        r"[а-яё]+\s+[а-яё]\.[а-яё]\.",
        r"[а-яё]\.[а-яё]\.\s+[а-яё]+",
        r"(?:директор|руководитель)\s*[а-яё]+\s*[а-яё]\.",
    ]

    target_runs = []

    last_paras = (
        document.main_paragraphs[-5:]
        if len(document.main_paragraphs) >= 5
        else document.main_paragraphs
    )
    for para in last_paras:
        if any(re.search(p, para.text, re.IGNORECASE) for p in patterns):
            target_runs = [run for run in para.runs if run.text.strip()]
            break

    if not target_runs and document.tables:
        for table in document.main_tables[:]:
            tbl_text = "\n".join(_get_table_text(document, table))
            if any(re.search(p, tbl_text, re.IGNORECASE) for p in patterns):
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            target_runs.extend(
                                [run for run in p.runs if run.text.strip()]
                            )
                break

    if not target_runs:
        document.errors.append(
            ValidationError(
                "SIGNATURE_BLOCK",
                "No head signature block found (expected 'Surname I.I.' or position near end).",
                Severity.CRITICAL,
                "End of main text",
            )
        )
        return

    for run in target_runs:
        size = _get_effective_font_size(document, run)
        if size is None or abs(size - 12.0) > 0.5:
            document.errors.append(
                ValidationError(
                    "SIGNATURE_FONT",
                    f"Signature font size is incorrect: {size}pt (expected 12pt)",
                    Severity.WARNING,
                    "Signature block",
                )
            )
            document.errors.append(
                ValidationError(
                    "SIGNATURE_FONT",
                    "Size and font in block can indicate blank table",
                    Severity.INFO,
                    "Signature block",
                )
            )
            return


def _check_executant_block(document):
    patterns = [
        r"(?:исполнитель|фамилия)\b.*?(?:доп|тел)\b",
        r"доп\.\s*\d{3,}",
        r"тел\.\s*[\d\-\(\)\s]{5,}",
        r"[а-яё]+\s+[а-яё]\.\s*[а-яё]\.\s*[\(]?(?:доп|тел)",
    ]

    target_runs = []

    last_paras = (
        document.main_paragraphs[-5:]
        if len(document.main_paragraphs) >= 5
        else document.main_paragraphs
    )
    for para in last_paras:
        if any(re.search(p, para.text, re.IGNORECASE) for p in patterns):
            target_runs = [run for run in para.runs if run.text.strip()]
            break

    if not target_runs and document.tables:
        for table in document.main_tables[:]:
            tbl_text = "\n".join(_get_table_text(document, table))
            if any(re.search(p, tbl_text, re.IGNORECASE) for p in patterns):
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            target_runs.extend(
                                [run for run in p.runs if run.text.strip()]
                            )
                break

    if not target_runs:
        document.errors.append(
            ValidationError(
                "EXECUTANT_BLOCK",
                "Executant block (ФИО + доп./тел.) not found at the end.",
                Severity.CRITICAL,
                "End of document",
            )
        )
        return

    for run in target_runs:
        size = _get_effective_font_size(document, run)
        if size is None or abs(size - 10.0) > 0.5:
            document.errors.append(
                ValidationError(
                    "EXECUTANT_FONT",
                    f"Executant font size is incorrect: {size}pt (expected 10pt)",
                    Severity.WARNING,
                    "Executant block",
                )
            )
            document.errors.append(
                ValidationError(
                    "EXECUTANT_FONT",
                    "Size and font in block can indicate blank table",
                    Severity.INFO,
                    "Tables",
                )
            )
            return


def _check_indents(doc):
    """
    Validates that first-line indents in main paragraphs match the standard (1.25 cm).
    Accounts for style inheritance, EMU unit conversion, and floating-point tolerance.
    """
    EXPECTED_INDENT_CM = 1.25
    TOLERANCE_CM = (
        0.05  # Allow small rounding differences from Word's internal precision
    )

    for para in doc.main_paragraphs[:-2]:
        # Skip empty paragraphs to avoid false positives
        if not para.text.strip():
            continue

        # Step 1: Try to get explicit first-line indent from paragraph formatting
        first_line = para.paragraph_format.first_line_indent

        # Step 2: If not explicitly set, check the paragraph style (Word often stores indents there)
        if first_line is None and para.style is not None:
            try:
                first_line = para.style.paragraph_format.first_line_indent
            except Exception:
                pass  # Style may not have indent defined; fall through to default handling

        # Step 3: Evaluate correctness and prepare display value
        is_incorrect = False
        display_value = "Not set (inherited)"

        if first_line is not None:
            # Convert EMU (English Metric Units) to centimeters for human-readable comparison
            indent_cm = first_line.cm
            display_value = f"{indent_cm:.2f} cm"
            # Use tolerance-based comparison to handle Word's internal rounding
            if abs(indent_cm - EXPECTED_INDENT_CM) > TOLERANCE_CM:
                is_incorrect = True
        else:
            # Standard requires explicit 1.25 cm indent; missing value is an issue
            is_incorrect = True

        if is_incorrect:
            # Truncate text for readability in the report
            preview = para.text.strip()
            preview = preview[:60] + "..." if len(preview) > 60 else preview

            doc.errors.append(
                ValidationError(
                    "PARAGRAPH_INDENTS",
                    f"First-line indent incorrect: {display_value} (expected {EXPECTED_INDENT_CM} cm)",
                    Severity.WARNING,
                    "Paragraph formatting",
                )
            )


def _check_orientation(doc):
    orientation = doc.doc.sections[0].orientation
    if orientation != WD_ORIENTATION.PORTRAIT:
        doc.errors.append(
            ValidationError(
                "PAGE_ORIENTATION",
                f"Page orientation incorrect: {orientation} (expected PORTRAIT)",
                Severity.WARNING,
                "Document body",
            )
        )
