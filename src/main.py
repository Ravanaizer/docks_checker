from typing import List

from docx import Document

from bookmarks import _check_bookmarks_presence
from check_appendix import (
    _check_appendix_font_name,
    _check_appendix_font_size,
    _check_appendix_format,
)
from check_body import (
    _check_command_word,
    _check_control_clause_position,
    _check_executant_block,
    _check_heading,
    _check_indents,
    _check_orientation,
    _check_preamble_structure,
    _check_reduction_position,
    _check_signature_block,
)
from check_fonts import (
    _check_body_font_name,
    _check_body_font_size,
    _check_list_font_name,
    _check_list_font_size,
)
from check_pages import _check_footer_empty_lines, _check_page_numbering
from check_spacing import _check_structural_spacing
from config import Severity, ValidationError
from tables import _check_table_fonts_name, _check_table_fonts_size, _get_main_tables
from text_normalization import (
    _clean_empty_paragraphs,
    _find_main_document_boundary,
    _normalize_text,
)


class DocumentArchitectureValidator:
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.errors: List[ValidationError] = []

        # 1. Инициализируем списки параграфов и таблиц сразу после загрузки
        self.paragraphs = self.doc.paragraphs
        self.tables = self.doc.tables

        # 2. Проверяем структуру отступов ДО очистки пустых параграфов
        # Нам нужно видеть пустые строки, чтобы проверить их количество
        _check_structural_spacing(self)

        # 3. Очищаем документ от пустых параграфов
        _clean_empty_paragraphs(self)

        # 4. ОБНОВЛЯЕМ self.paragraphs после очистки.
        # python-docx возвращает новый список при обращении к свойству.
        # Это критично, так как _find_main_document_boundary вернет индекс
        # в очищенном документе, и нам нужно нарезать именно очищенный список.
        self.paragraphs = self.doc.paragraphs

        # 5. Поиск границы основного текста и приложений
        self.main_doc_end_idx = _find_main_document_boundary(self)

        # 6. Разделение частей документа
        self.main_paragraphs = self.paragraphs[: self.main_doc_end_idx]
        self.appendix_paragraphs = self.paragraphs[self.main_doc_end_idx :]
        self.main_tables = []
        _get_main_tables(self)

        # 7. Нормализация текста для проверок содержимого
        self.main_text = _normalize_text(
            "\n".join([p.text.strip() for p in self.main_paragraphs if p.text.strip()])
        )
        self.full_text = _normalize_text(
            "\n".join([p.text.strip() for p in self.paragraphs if p.text.strip()])
        )

    def validate(self) -> List[ValidationError]:
        _check_signature_block(self)
        _check_executant_block(self)
        _check_heading(self)
        _check_body_font_size(self)
        _check_preamble_structure(self)
        _check_command_word(self)
        _check_control_clause_position(self)
        _check_table_fonts_size(self)
        _check_table_fonts_name(self)
        _check_reduction_position(self)
        _check_list_font_size(self)
        _check_body_font_name(self)
        _check_list_font_name(self)
        _check_indents(self)
        _check_appendix_format(self)
        _check_appendix_font_name(self)
        _check_appendix_font_size(self)
        _check_orientation(self)
        _check_page_numbering(self)
        _check_footer_empty_lines(self)
        _check_bookmarks_presence(self, ["data", "nomer"])
        return self.errors

    def print_report(self) -> None:
        if not self.errors:
            print("Document architecture meets requirements.")
            return

        critical = [e for e in self.errors if e.severity == Severity.CRITICAL]
        warnings = [e for e in self.errors if e.severity == Severity.WARNING]
        info = [e for e in self.errors if e.severity == Severity.INFO]

        print("\n" + "=" * 70)
        print("DOCUMENT ARCHITECTURE VALIDATION REPORT")
        print("=" * 70)
        print(f"File: {self.filepath}")
        print(f"Total findings: {len(self.errors)}\n")

        if critical:
            print("CRITICAL ERRORS:")
            for err in critical:
                print(f"  {err.severity.value} | {err.rule}")
                print(f"    {err.message}")
                print()

        if warnings:
            print("WARNINGS:")
            for err in warnings:
                print(f"  {err.severity.value} | {err.rule}")
                print(f"    {err.message}")
                print()

        print("=" * 70)

        if info:
            print("INFO:")
            for err in info:
                print(f"  {err.severity.value} | {err.rule}")
                print(f"    {err.message}")
                print()

        print("=" * 70)


if __name__ == "__main__":
    # for sdk in range(1, 15):
    #     validator = DocumentArchitectureValidator(
    #         f"../tests/{sdk}.docx"
    #     )  # "standart.docx")  #
    #     validator.validate()
    #     validator.print_report()

    validator = DocumentArchitectureValidator(
        "standart.docx"
    )  # f"../tests/13.docx")  #
    validator.validate()
    validator.print_report()
