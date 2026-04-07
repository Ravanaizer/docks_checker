import re

from docx.oxml.ns import qn


def _normalize_text(text: str) -> str:
    """Нормализация текста: замена множественных пробелов/табуляций на один пробел."""
    if not isinstance(text, str):
        return ""
    return re.sub(r"\s+", " ", text).strip()


def _clean_empty_paragraphs(document) -> None:
    """Удаление пустых параграфов из документа."""
    for p in reversed(document.doc.paragraphs):
        if not p.text.strip():
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)


def _find_main_document_boundary(document) -> int:
    """Поиск границы между основным текстом и приложениями."""
    appendix_markers = [
        r"^приложение\s*[№№]?\s*[а-яё1-9]",
        r"^прил\.\s*",
    ]

    para_idx = 0
    for child in document.doc.element.body:
        if child.tag == qn("w:p"):
            if para_idx < len(document.paragraphs):
                para = document.paragraphs[para_idx]
                text = _normalize_text(para.text.strip().lower())
                if any(re.search(pattern, text) for pattern in appendix_markers):
                    return para_idx
                para_idx += 1
        elif child.tag == qn("w:tbl"):
            for row in child.iter(qn("w:tr")):
                for cell in row.iter(qn("w:tc")):
                    for p in cell.iter(qn("w:p")):
                        cell_text = " ".join(p.itertext()).strip().lower()
                        cell_text = _normalize_text(cell_text)
                        if any(
                            re.search(pattern, cell_text)
                            for pattern in appendix_markers
                        ):
                            return para_idx
    return len(document.paragraphs)
